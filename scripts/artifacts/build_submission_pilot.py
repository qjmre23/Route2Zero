import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
TMP = ROOT / "tmp" / "route2zero-submission"
BANNER = TMP / "hackathon-template-banner.png"
OUT = ROOT / "output" / "submission" / "Route2Zero_Team_Larpers_Pilot_Plan.docx"
GITHUB = "https://github.com/qjmre23/Route2Zero"
HOST = "https://route2zero.netlify.app/"
MANIFEST = json.loads((ROOT / "data" / "processed" / "build_manifest.json").read_text(encoding="utf-8"))
FEASIBILITY = json.loads((ROOT / "data" / "processed" / "feasibility_cost_scenarios.json").read_text(encoding="utf-8"))
BUILD_ID = MANIFEST["build_id"]
SCENARIO_ID = MANIFEST["default_scenario_id"]
PORTFOLIO_ID = MANIFEST["default_portfolio_scenario_id"]
MODEL_VERSION = MANIFEST["model_versions"]["service_intensity"]

PALE_BLUE = "CFE2F3"
DEEP_BLUE = "0B2E59"
INK = "111111"
MUTED = "4B5563"
LINK = "075A91"
SOFT = "F5F8FC"


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trpr.append(el)


def no_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, color="000000", size="9"):
    tcpr = cell._tc.get_or_add_tcPr()
    group = tcpr.first_child_found_in("w:tcBorders")
    if group is None:
        group = OxmlElement("w:tcBorders")
        tcpr.append(group)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn("w:" + edge)
        el = group.find(tag)
        if el is None:
            el = OxmlElement("w:" + edge)
            group.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def cell_margins(cell, top=72, start=100, bottom=72, end=100):
    tcpr = cell._tc.get_or_add_tcPr()
    group = tcpr.first_child_found_in("w:tcMar")
    if group is None:
        group = OxmlElement("w:tcMar")
        tcpr.append(group)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = group.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            group.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def cell_width(cell, inches):
    cell.width = Inches(inches)
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(round(inches * 1440)))
    tcw.set(qn("w:type"), "dxa")


def table_geometry(table_obj, widths, indent=100):
    values = [round(v * 1440) for v in widths]
    total = sum(values)
    tblpr = table_obj._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table_obj._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in values:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table_obj.rows:
        for idx, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(values[idx]))
            tcw.set(qn("w:type"), "dxa")


def font(run, size=9, bold=False, italic=False, color=INK):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def hyperlink(paragraph, label, url, size=8.3):
    rid = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    rpr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK)
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(round(size * 2)))
    rpr.append(sz)
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    link.append(run)
    paragraph._p.append(link)


def page_field(paragraph):
    r = paragraph.add_run("Page ")
    font(r, 7, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    r = paragraph.add_run(" of 5")
    font(r, 7, color=MUTED)


def page_header(doc, label, force_page=False):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = force_page
    p.paragraph_format.left_indent = Inches(-0.78)
    p.paragraph_format.right_indent = Inches(-0.78)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(BANNER), width=Cm(21.0))
    inline = run._r.xpath(".//wp:inline")[0]
    prop = inline.xpath("./wp:docPr")[0]
    prop.set("descr", "AI x City Climate Action Hackathon 2026 partner banner")
    prop.set("title", "AI x City Climate Action Hackathon 2026")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    hyperlink(p, f"GitHub: {GITHUB}", GITHUB)
    r = p.add_run("   |   ")
    font(r, 8.3, color=MUTED)
    hyperlink(p, f"Live dashboard: {HOST}", HOST)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(label)
    font(r, 7.2, bold=True, color=MUTED)


def title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), 19, bold=True)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.keep_with_next = True
        font(p.add_run(subtitle), 10.2, color=DEEP_BLUE)


def section(doc, text, after=3):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text.upper()), 11.5, bold=True, color=DEEP_BLUE)


def body(doc, text, size=9, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.06
    font(p.add_run(text), size)


def bullets(doc, items, size=8.7, after=1):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.0
        font(p.add_run(item), size)


def callout(doc, heading, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    cell_width(cell, 6.72)
    shade(cell, SOFT)
    borders(cell, "9FBAD0", "8")
    cell_margins(cell, 110, 140, 110, 140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    font(p.add_run(heading), 9.4, bold=True, color=DEEP_BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    font(p.add_run(text), 9)
    table_geometry(table, [6.72], indent=140)


def table(doc, headers, rows, widths, size=8, header_size=8.4, stripe=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    repeat_header(t.rows[0])
    no_split(t.rows[0])
    for i, text in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell_width(cell, widths[i])
        shade(cell, PALE_BLUE)
        borders(cell)
        cell_margins(cell, 82, 100, 82, 100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(text), header_size)
    for row_i, values in enumerate(rows):
        cells = t.add_row().cells
        no_split(t.rows[-1])
        for i, value in enumerate(values):
            cell = cells[i]
            cell_width(cell, widths[i])
            borders(cell)
            cell_margins(cell, 68, 100, 68, 100)
            if stripe and row_i % 2:
                shade(cell, "F8FAFC")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 0.98
            font(p.add_run(value), size)
    table_geometry(t, widths)
    return t


doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Inches(0.15)
sec.bottom_margin = Inches(0.42)
sec.left_margin = Inches(0.78)
sec.right_margin = Inches(0.78)
sec.header_distance = Inches(0.1)
sec.footer_distance = Inches(0.18)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(9)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(3)

title_style = doc.styles["Title"]
title_style.font.name = "Arial"
title_style.font.size = Pt(19)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor.from_string(INK)
heading_style = doc.styles["Heading 1"]
heading_style.font.name = "Arial"
heading_style.font.size = Pt(11.5)
heading_style.font.bold = True
heading_style.font.color.rgb = RGBColor.from_string(DEEP_BLUE)

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run(f"Team Larpers | Route2Zero | {BUILD_ID} | {SCENARIO_ID} | {PORTFOLIO_ID} | "), 5.8, color=MUTED)
page_field(footer)

# Page 1 — pilot purpose and outcome.
page_header(doc, "TEAM LARPERS / ROUTE2ZERO / PILOT PLAN")
title(doc, "Team Larpers Pilot Plan", "Route2Zero Six-Month City Pilot")
callout(
    doc,
    "City pilot question",
    "Can Route2Zero produce a city-owned, evidence-backed e-jeepney Phase-1 portfolio whose climate, equity, infrastructure, operator, and uncertainty assumptions have been tested against real local evidence?",
)
section(doc, "Pilot purpose")
body(doc, "Convert Route2Zero 2.1 from an evidence-based prototype into a locally validated planning workflow that cities can inspect, challenge, refresh, and own.")
section(doc, "Pilot scope")
bullets(
    doc,
    [
        "Begin with one lead LGU; add one comparison LGU only after Gate 1 and a signed second-city pilot agreement.",
        "Validate a manageable corridor sample, one flagship corridor, and a broader Phase-1 shortlist.",
        "Engage relevant operators/cooperatives, a utility or energy counterpart, and rider/community representatives where appropriate.",
        "Test ML service estimates, climate assumptions, route geometry, equity evidence, charging evidence, operator readiness, policy scenarios, robustness, and portfolio constraints.",
    ],
)
section(doc, "Six-month outcome")
table(
    doc,
    ["Outcome area", "Deliverable"],
    [
        ("Decision", "City-owned Phase-1 portfolio or shortlist with named next actions."),
        ("Evidence", "Validated route cards, source registry, validation ledgers, and unresolved evidence owners."),
        ("Models", "Calibrated service-intensity model, climate scenarios, model card, and evidence-confidence rules."),
        ("Operations", "Scenario library, decision pack, refresh/retraining runbook, and reproducibility handover."),
    ],
    [1.35, 5.37],
    size=8.5,
    header_size=8.7,
    stripe=True,
)
section(doc, "Pilot pathway")
flow = doc.add_table(rows=1, cols=5)
flow.alignment = WD_TABLE_ALIGNMENT.CENTER
flow.autofit = False
repeat_header(flow.rows[0])
for i, label in enumerate(("PROTOTYPE", "VALIDATE", "CALIBRATE", "CO-DESIGN", "HANDOVER")):
    cell = flow.cell(0, i)
    cell_width(cell, 1.34)
    shade(cell, PALE_BLUE if i % 2 == 0 else SOFT)
    borders(cell, "9FBAD0", "8")
    cell_margins(cell, 100, 40, 100, 40)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(label), 7.8, bold=True, color=DEEP_BLUE)
table_geometry(flow, [1.34, 1.34, 1.34, 1.34, 1.34], indent=40)

# Page 2 — official-style timeline table.
page_header(doc, "SIX MONTHS / THREE EVIDENCE GATES", True)
title(doc, "Six months, three evidence gates")
timeline = [
    ("Month 1", "Baseline + model protocol", "Confirm decision owners and route sample; freeze source registry/build; approve validation, ML evaluation, privacy, and climate protocols.", "Pilot charter; baseline build; route sample; validation protocol; risk register."),
    ("Month 2", "Data refresh + ground truth", "Extend the hackathon's 20-route OSM validation seed with field observation; validate geometry and service; collect operator, charging/depot, and local equity/accessibility evidence.", "Expanded validation, operator, and charging ledgers; ground-truth sample; source refresh."),
    ("Gate 1", "Evidence readiness review", "Accept evidence plan and model-validation protocol; confirm minimum data coverage; resolve governance blockers.", "Approved evidence gate or documented corrective actions."),
    ("Month 3", "Model + route validation", "Compare ML estimates with refreshed observations; recalculate metrics; review anomalies; tune geometry/evidence rules and equity variables.", "Pilot model evaluation; model-card update; revised evidence grades; validated flagship case."),
    ("Month 4", "Climate, charging + operator pre-feasibility", "Calibrate diesel/electric assumptions; validate energy scenarios; request utility evidence; assess sites, operators, and verified constraints.", "Calibrated climate scenario; charging/operator cards; selection-ready constraints."),
    ("Gate 2", "Quality + route-set review", "Accept pilot route set and model/evidence quality; hold routes with unresolved critical issues or mark them evidence-limited.", "Approved route set with explicit holds and evidence limits."),
    ("Month 5", "Scenario co-design + human factors", "Run policy workshops; review sensitivity; agree portfolio constraints; test Priority vs Evidence vs Stability; review assistant and accessibility.", "Approved scenarios; draft portfolio; usability and governance findings."),
    ("Month 6", "Decision pack + handover", "Finalize portfolio, evidence queue, model card, source registry, refresh process, issue log, analyst training, and task ownership.", "Accepted decision pack; final build; reproducibility package; city handover."),
    ("Gate 3", "Acceptance", "City decision pack and reproducibility handover accepted by named decision owners.", "Go-forward decision and accountable next-step owners."),
]
table(doc, ["Timeline", "Activity / Milestone", "Core activity", "Outcome"], timeline, [0.72, 1.42, 2.75, 1.83], size=7.15, header_size=7.7, stripe=True)
body(doc, f"Baseline traceability: build {BUILD_ID} | scenario {SCENARIO_ID} | service model {MODEL_VERSION}", 7.4, 0)

# Page 3 — validation and responsible AI.
page_header(doc, "VALIDATION / SUCCESS / RESPONSIBLE AI", True)
title(doc, "The pilot succeeds when the evidence and decision improve")
section(doc, "Success metrics", 2)
table(
    doc,
    ["Measure", "Pilot success evidence"],
    [
        ("Coverage", "% pilot routes ground-checked; % shortlisted routes with complete evidence cards; critical gaps resolved."),
        ("Model", "MAE/RMSE vs refreshed observations, improvement over baseline, and documented route-level residual review."),
        ("AI task value", "Compare AI-assisted and analyst-only evidence review using time-to-first brief, missed-evidence rate, factual errors, and unsupported claims."),
        ("Decision quality", "Evidence-grade improvement; ranking/portfolio stability before vs after validation; routes moved from evidence-limited to decision-ready."),
        ("Adoption", "LGU/operator/utility interactions; planner task completion; understanding of Priority, Evidence, and Stability."),
        ("Handover", "One accepted city decision pack and one completed reproducibility/refresh handover."),
    ],
    [1.25, 5.47],
    size=7.8,
    header_size=8.3,
    stripe=True,
)
section(doc, "Validation methods", 2)
body(doc, "Field observations; operator interviews; LGU review; utility evidence requests; data reconciliation; grouped holdout testing; scenario workshops; usability/accessibility testing; stakeholder review of the flagship corridor.", 8.3)
section(doc, "Responsible AI controls", 2)
bullets(
    doc,
    [
        "No resident-level personal data are required for core analytics; stakeholder data are consent-based.",
        "No hidden LLM score changes and no automatic funding or procurement decisions.",
        "No individual poverty or informal-status inference from population density.",
        "Stakeholders can challenge or replace evidence; an AI-disabled deterministic fallback remains available.",
    ],
    8.0,
    0,
)
section(doc, "Risk register", 2)
table(
    doc,
    ["Risk", "Control / response"],
    [
        ("Historic route data", "Use current-validation ledger and field checks; retain historic-only status until verified."),
        ("Weak ML performance", "Prefer observed evidence; mark model experimental; compare with baseline."),
        ("Utility capacity unavailable", "Show mapped proximity only; keep capacity NOT VERIFIED; request evidence."),
        ("Operator data unavailable", "Keep neutral prior explicit and raise validation priority; a named desk reference is not readiness evidence."),
        ("Cost/financing data unverified", "Treat feasibility figures as PROXY/SCENARIO until fleet and tariff data are confirmed; never present them as a budget."),
        ("Sparse geometry / equity data", "Use reliability/evidence grades; validate traces; avoid unsupported settlement claims."),
        ("AI/API unavailable", "Serve structured deterministic fallback; scores, scenarios, and portfolio are unaffected."),
        ("Constraint conflict", "Show infeasibility and binding constraints; never fabricate a portfolio."),
    ],
    [1.55, 5.17],
    size=7.5,
    header_size=8.1,
    stripe=True,
)

# Page 4 — financial feasibility and licensing.
page_header(doc, "FINANCIAL FEASIBILITY / DATA LICENSING", True)
title(doc, "Validate the economics with the same discipline as the climate case")
callout(
    doc,
    "Starting-point scenario — not a budget",
    "The Phase-1 shortlist currently implies about 1,943 vehicles, 102 chargers, and PHP 4.91 billion in vehicle-plus-charger hardware under explicit PROXY assumptions. Financing terms remain MISSING, and excluded costs are not silently absorbed.",
)
section(doc, "Financial feasibility approach", 2)
table(
    doc,
    ["Input", "Current status", "Pilot validation needed"],
    [
        ("Fleet size", "1,943 vehicles | PROXY", "Confirm active fleet, duty cycle, service plan, spare ratio, and replacement phasing by corridor."),
        ("Vehicle cost", "PHP 2.5M/unit | PROXY", "Obtain supplier-neutral market evidence and define vehicle specification; do not treat the DOE planning value as a jeepney quote."),
        ("Charging", "102 stations | PROXY", "Validate charger power, depot layout, operating windows, utility capacity, interconnection, and civil works."),
        ("Financing", "MISSING", "Confirm tariff structure, equity/debt source, grants or subsidy, repayment terms, insurance, and ownership model."),
    ],
    [1.15, 1.65, 3.92],
    size=7.7,
    header_size=8.0,
    stripe=True,
)
body(doc, "Excluded from the current capital proxy: land or depot acquisition, civil works, distribution-system upgrades, interconnection fees, battery replacement, taxes, insurance, financing costs, and operations and maintenance.", 8.0)
section(doc, "Data licensing & attribution", 2)
bullets(
    doc,
    [
        "Project code is MIT-licensed; source data retain their own licenses and conditions.",
        "OpenStreetMap route and infrastructure data require attribution to OpenStreetMap contributors and are used under ODbL. Attribution must remain visible in maps, exports, decision packs, and future field-data updates.",
        "The source registry records URLs, retrieval dates, license notes, checksums, and claim scope. Pilot owners review these before every release.",
        "New local evidence is accepted only with an accountable source, observation date, permission basis where relevant, and a field-level claim status.",
    ],
    8.0,
    1,
)
section(doc, "Submission head start", 2)
body(doc, "The evidence set includes 20 dated OSM route matches, 20 observed member-way geometries, and 22 usable source geometries. These records reduce the initial desk-review burden but do not prove active operations or franchise authority.", 8.2)
body(doc, f"Baseline traceability: build {BUILD_ID} | scenario {SCENARIO_ID} | portfolio {PORTFOLIO_ID}", 7.4, 0)

# Page 5 — team structure and handover.
page_header(doc, "TEAM STRUCTURE / PARTNERS / HANDOVER", True)
title(doc, "Named owners from data collection to city handover")
section(doc, "Team structure", 2)
team = [
    ("John Marwin Ebona", "Product & Pilot Lead", "City coordination, decision question, gates, integration, decision pack, handover."),
    ("Isaac Marcus", "Engineering & Deployment Lead", "Dashboard, scenarios, Netlify, Mapbox, reliability, and deterministic fallbacks."),
    ("Andrei Dela Cruz", "Policy & Responsible AI Lead", "Policy controls, decision rights, workshops, risk controls, and responsible AI."),
    ("Russel Mendez", "Climate & Evidence Lead", "Climate and energy method, source review, evidence confidence, and documentation."),
    ("Joaquin Sarmiento", "Data & ML Lead", "Source registry, feature store, service model, metrics, manifest, and model refresh."),
    ("Prince Marl Mirasol", "QA & Documentation Lead", "Testing, accessibility, artifact QA, claim checks, and issue tracking."),
    ("Carl Nueva", "Geospatial & Infrastructure Analyst", "Geometry review, spatial joins, equity layers, power evidence, and map QA."),
    ("JOSEPH CLARENCE PARAYAOAN", "Field Validation & Partnerships Lead", "Field plans, route checks, operator and partner coordination, and evidence handoff."),
]
table(doc, ["Name", "Role", "Relevant background / experience / skills for this role"], team, [1.92, 1.80, 3.00], size=7.1, header_size=7.35, stripe=True)
section(doc, "Partner responsibilities", 2)
table(
    doc,
    ["Partner", "Pilot responsibility"],
    [
        ("LGU", "Define decision context, validate evidence, join scenario workshops, review outputs, accept handover."),
        ("Operator / cooperative", "Validate operations, fleet/depot constraints, readiness, and feasibility."),
        ("Utility / energy", "Review charging assumptions and flag formal capacity/interconnection studies."),
        ("Riders / community", "Validate access concerns and challenge equity assumptions where relevant."),
    ],
    [1.45, 5.27],
    size=6.9,
    header_size=7.5,
    stripe=True,
)
section(doc, "Handover and scale", 2)
body(doc, "Handover: final build, model card, source registry, field-observation template, validation ledgers, scenario library, decision pack, issue register, scheduled source-freshness check, refresh/retraining runbook, governance notes, and a city owner for each unresolved issue. Scale requires local adapters, recalibration, and validation; the Metro Manila model is not copied unchanged.", 7.3, 0)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "Team Larpers - Route2Zero Six-Month City Pilot Plan"
doc.core_properties.author = "Team Larpers"
doc.core_properties.subject = "AI x City Climate Action Hackathon 2026 Pilot Plan"
doc.core_properties.keywords = "Route2Zero, e-jeepney, climate action, urban transportation, pilot plan"
doc.save(OUT)
print(f"Created {OUT}")
